import os
import logging
from typing import Dict, List, Any, Optional

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing for various file formats"""

    def __init__(self):
        # Configure Tesseract if needed
        # pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'
        pass

    def extract_text(self, file_path: str) -> str:
        """Extract text from document based on file extension"""
        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return f"Error processing document: {str(e)}"

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods including OCR"""
        try:
            # Try PyMuPDF first
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_content = []
                for page in doc:
                    text_content.append(page.get_text())
                doc.close()
                if text_content and any(text.strip() for text in text_content):
                    extracted_text = "\n".join(text_content)
                    # Check if we got meaningful text (more than just whitespace/special chars)
                    if len(extracted_text.strip()) > 50:
                        return extracted_text
                    else:
                        logger.info("PyMuPDF extracted minimal text, trying OCR fallback")
            except ImportError:
                logger.warning("PyMuPDF not available, trying pdfplumber")
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {str(e)}")

            # Try pdfplumber
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text_content = []
                    for page in pdf.pages:
                        text_content.append(page.extract_text() or "")
                    if text_content and any(text.strip() for text in text_content):
                        extracted_text = "\n".join(text_content)
                        if len(extracted_text.strip()) > 50:
                            return extracted_text
                        else:
                            logger.info("pdfplumber extracted minimal text, trying OCR fallback")
            except ImportError:
                logger.warning("pdfplumber not available")
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {str(e)}")

            # If text extraction failed or returned minimal text, try OCR
            logger.info("Attempting OCR extraction as fallback")
            ocr_text = self._extract_with_ocr(file_path)
            if ocr_text and len(ocr_text.strip()) > 50:
                return ocr_text

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")

        return "PDF text extraction not available - install PyMuPDF or pdfplumber"

    def _extract_with_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR (for image-based PDFs)"""
        try:
            import fitz  # PyMuPDF for PDF to image conversion
            import pytesseract
            from PIL import Image
            import io
            
            logger.info(f"Starting OCR extraction for {file_path}")
            
            doc = fitz.open(file_path)
            text_content = []
            
            # Process each page (limit to first 5 pages for performance)
            max_pages = min(5, doc.page_count)
            
            for page_num in range(max_pages):
                try:
                    page = doc.load_page(page_num)
                    
                    # Convert page to image at higher resolution for better OCR
                    mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    image = Image.open(io.BytesIO(img_data))
                    
                    # Perform OCR
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    
                    if page_text.strip():
                        text_content.append(page_text)
                        logger.info(f"OCR extracted {len(page_text)} characters from page {page_num + 1}")
                    
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num + 1}: {str(e)}")
                    continue
            
            doc.close()
            
            if text_content:
                full_text = "\n\n".join(text_content)
                logger.info(f"OCR extraction completed: {len(full_text)} total characters")
                return full_text
            else:
                logger.warning("OCR extraction returned no text")
                return ""
                
        except ImportError as e:
            logger.warning(f"OCR libraries not available: {str(e)}")
            return ""
        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            return ""

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            # Try python-docx first
            try:
                from docx import Document
                doc = Document(file_path)
                text_content = []
                for paragraph in doc.paragraphs:
                    text_content.append(paragraph.text)
                return "\n".join(text_content)
            except ImportError:
                logger.warning("python-docx not available, trying docx2txt")
            except Exception as e:
                logger.warning(f"python-docx extraction failed: {str(e)}")

            # Try docx2txt fallback
            try:
                import docx2txt
                return docx2txt.process(file_path)
            except ImportError:
                logger.warning("docx2txt not available")
            except Exception as e:
                logger.error(f"docx2txt extraction failed: {str(e)}")

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")

        return "DOCX text extraction not available - install python-docx or docx2txt"

    def analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze document structure and formatting"""
        structure_info = {
            "font_analysis": {"unique_fonts": 0, "font_list": []},
            "layout_analysis": {"consistent_fonts": True},
            "page_count": 1,
            "image_count": 0
        }

        try:
            if file_path.lower().endswith('.pdf'):
                structure_info.update(self._analyze_pdf_structure(file_path))
            elif file_path.lower().endswith('.docx'):
                structure_info.update(self._analyze_docx_structure(file_path))
        except Exception as e:
            logger.error(f"Structure analysis failed: {str(e)}")

        return structure_info

    def _analyze_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze PDF structure"""
        try:
            import fitz
            doc = fitz.open(file_path)
            fonts = set()
            pages_info = []

            for page_num in range(min(3, doc.page_count)):  # Limit analysis
                page = doc.load_page(page_num)
                page_fonts = set()

                try:
                    for block in page.get_text("dict")["blocks"]:
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    font_info = f"{span['font']}:{span['size']}"
                                    page_fonts.add(font_info)
                except:
                    pass  # Skip if text extraction fails

                fonts.update(page_fonts)
                pages_info.append({
                    "page": page_num + 1,
                    "fonts": list(page_fonts),
                    "text_length": len(page.get_text())
                })

            doc.close()

            return {
                "font_analysis": {
                    "unique_fonts": len(fonts),
                    "font_list": list(fonts),
                    "pages_info": pages_info
                },
                "page_count": len(pages_info),
                "layout_analysis": {
                    "consistent_fonts": len(set(tuple(page["fonts"]) for page in pages_info)) <= 1
                }
            }
        except Exception as e:
            logger.error(f"PDF structure analysis failed: {str(e)}")
            return {}

    def _analyze_docx_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze DOCX structure"""
        try:
            from docx import Document
            doc = Document(file_path)
            fonts = set()

            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name:
                        fonts.add(f"{run.font.name}:{run.font.size}")

            return {
                "font_analysis": {
                    "unique_fonts": len(fonts),
                    "font_list": list(fonts)
                },
                "page_count": 1,
                "layout_analysis": {
                    "consistent_fonts": len(fonts) <= 3
                }
            }
        except Exception as e:
            logger.error(f"DOCX structure analysis failed: {str(e)}")
            return {}
